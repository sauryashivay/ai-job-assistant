from io import BytesIO
from pathlib import Path

import fitz
from docx import Document


SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


def extract_pdf_text(file_bytes: bytes) -> str:
    document = fitz.open(
        stream=file_bytes,
        filetype="pdf",
    )

    pages: list[str] = []

    try:
        for page in document:
            text = page.get_text("text")

            if text.strip():
                pages.append(text)
    finally:
        document.close()

    return "\n".join(pages).strip()


def extract_docx_text(file_bytes: bytes) -> str:
    document = Document(BytesIO(file_bytes))

    paragraphs = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]

    for table in document.tables:
        for row in table.rows:
            values = [
                cell.text.strip()
                for cell in row.cells
                if cell.text.strip()
            ]

            if values:
                paragraphs.append(" | ".join(values))

    return "\n".join(paragraphs).strip()


def extract_resume_text(
    filename: str,
    file_bytes: bytes,
) -> str:
    extension = Path(filename).suffix.lower()

    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            "Unsupported file type. Upload a PDF or DOCX resume."
        )

    if extension == ".pdf":
        text = extract_pdf_text(file_bytes)
    else:
        text = extract_docx_text(file_bytes)

    if not text.strip():
        raise ValueError(
            "No readable text was found in the uploaded resume."
        )

    return text